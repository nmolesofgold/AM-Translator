import io
import os
import re
import time
import shutil
import hashlib
import unicodedata

import streamlit as st
import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
from docx import Document
from openpyxl import load_workbook
from PIL import Image
import pytesseract


# -----------------------------
# Page Configuration
# -----------------------------
st.set_page_config(
    page_title="Universal Document Translator",
    page_icon="🌐",
    layout="wide",
    initial_sidebar_state="collapsed",
)


# -----------------------------
# Session State
# -----------------------------
DEFAULT_SESSION_STATE = {
    "processed_output_data": None,
    "processed_output_name": "",
    "translated_text_result": "",
    "debug_log": [],
    "translation_cache": {},
    "preserve_formula_tokens": True,
}

for key, value in DEFAULT_SESSION_STATE.items():
    if key not in st.session_state:
        st.session_state[key] = value


# -----------------------------
# Language Mapping
# -----------------------------
SOURCE_LANGUAGES = {
    "Auto Detect": "auto",
    "German": "de",
    "English": "en",
    "Spanish": "es",
    "French": "fr",
    "Italian": "it",
    "Portuguese": "pt",
    "Chinese Simplified": "zh-CN",
    "Japanese": "ja",
    "Korean": "ko",
    "Russian": "ru",
    "Arabic": "ar",
    "Hindi": "hi",
    "Dutch": "nl",
    "Polish": "pl",
    "Turkish": "tr",
    "Swedish": "sv",
}

TARGET_LANGUAGES = {
    "German": "de",
    "English": "en",
    "Spanish": "es",
    "French": "fr",
    "Italian": "it",
    "Portuguese": "pt",
    "Chinese Simplified": "zh-CN",
    "Japanese": "ja",
    "Korean": "ko",
    "Russian": "ru",
    "Arabic": "ar",
    "Hindi": "hi",
    "Dutch": "nl",
    "Polish": "pl",
    "Turkish": "tr",
    "Swedish": "sv",
}

# Tesseract uses different language codes than Google Translate.
TESSERACT_LANGUAGES = {
    "auto": "eng",       # Tesseract does not truly auto-detect.
    "en": "eng",
    "de": "deu",
    "es": "spa",
    "fr": "fra",
    "it": "ita",
    "pt": "por",
    "zh-CN": "chi_sim",
    "ja": "jpn",
    "ko": "kor",
    "ru": "rus",
    "ar": "ara",
    "hi": "hin",
    "nl": "nld",
    "pl": "pol",
    "tr": "tur",
    "sv": "swe",
}


# -----------------------------
# Unicode Font Support
# -----------------------------
COMMON_UNICODE_FONTS = [
    # Streamlit Cloud / Debian / Ubuntu
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",

    # Noto fonts
    "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",

    # macOS fallback
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/Library/Fonts/Arial Unicode.ttf",
    "/Library/Fonts/Arial.ttf",

    # Windows fallback
    r"C:\Windows\Fonts\arial.ttf",
    r"C:\Windows\Fonts\segoeui.ttf",
]


def find_unicode_font():
    for font_path in COMMON_UNICODE_FONTS:
        if os.path.exists(font_path):
            return font_path
    return None


UNICODE_FONT_PATH = find_unicode_font()
CUSTOM_FONT_NAME = "customunicodefont"


# -----------------------------
# Header
# -----------------------------
st.title("🌐 Universal Document Translator")
st.caption("Translate PDFs, Word docs, Excel sheets, and scanned PDFs with optional OCR.")


# -----------------------------
# Sidebar
# -----------------------------
with st.sidebar:
    st.header("Settings")

    source_language = st.selectbox(
        "Source Language",
        options=list(SOURCE_LANGUAGES.keys()),
        index=0,
    )

    target_language = st.selectbox(
        "Target Language",
        options=list(TARGET_LANGUAGES.keys()),
        index=1,  # Default English
    )

    src_code = SOURCE_LANGUAGES[source_language]
    tgt_code = TARGET_LANGUAGES[target_language]

    st.divider()

    debug_mode = st.checkbox("Enable Error Logging", value=False)

    st.subheader("PDF/Text Safety")

    preserve_formula_tokens = st.checkbox(
        "Preserve formula/scientific tokens",
        value=True,
        help=(
            "Helps prevent translation from damaging tokens like CO₂, H₂O, x², 10⁻³, etc. "
            "This cannot fix characters that the PDF extractor or OCR already read incorrectly."
        ),
    )

    st.session_state["preserve_formula_tokens"] = preserve_formula_tokens

    st.subheader("OCR Settings")

    enable_ocr = st.checkbox(
        "Enable OCR for scanned PDFs",
        value=True,
        help="Uses free local Tesseract OCR for image-only/scanned PDF pages.",
    )

    ocr_output_mode = st.selectbox(
        "OCR Output Mode",
        options=[
            "Append translated OCR pages",
            "Overlay translated text",
        ],
        index=0,
        help=(
            "Append mode is more reliable. "
            "Overlay mode covers the scanned page with translated text and may not preserve layout."
        ),
    )

    ocr_zoom = st.slider(
        "OCR Image Quality",
        min_value=1.5,
        max_value=4.0,
        value=2.5,
        step=0.5,
        help="Higher values can improve OCR accuracy but use more memory.",
    )

    if enable_ocr and src_code == "auto":
        st.warning(
            "OCR works best when you choose the real source language instead of Auto Detect. "
            "For German scanned PDFs, choose Source Language = German."
        )

    st.divider()
    st.subheader("Server Status")

    tesseract_path = shutil.which("tesseract")

    if tesseract_path:
        st.success(f"Tesseract found: {tesseract_path}")
    else:
        st.error("Tesseract not found. Check packages.txt and redeploy.")

    if UNICODE_FONT_PATH:
        st.success(f"Unicode font found: {os.path.basename(UNICODE_FONT_PATH)}")
    else:
        st.warning("No Unicode font found. Some characters may become ? in PDFs.")

    if debug_mode and tesseract_path:
        try:
            installed_ocr_langs = pytesseract.get_languages(config="")
            st.caption("OCR languages installed:")
            st.code(", ".join(installed_ocr_langs))
        except Exception as exc:
            st.caption(f"Could not list OCR languages: {exc}")


# -----------------------------
# Text Cleanup / Normalization
# -----------------------------
def normalize_ocr_and_pdf_text(text: str) -> str:
    """
    Normalize text safely.

    Important:
    - Uses NFC, not NFKC.
    - NFKC can convert superscripts/subscripts into regular characters.
      For example, ² may become 2. We do not want that.
    """
    if not text:
        return ""

    cleaned = unicodedata.normalize("NFC", text)

    replacements = {
        "\u00A0": " ",  # non-breaking space
        "``": "“",
        "''": "”",
        "´´": "“",
        "…": "...",

        # Common OCR/PDF confusions. Conservative replacements only.
        "„ ": "„",
        " “": " “",
    }

    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)

    return cleaned


def sanitize_text(text: str) -> str:
    if not text:
        return ""

    replacements = {
        "\u00A0": " ",
        "\u200B": "",
        "\uFEFF": "",
        "\u200C": "",
        "\u200D": "",
    }

    cleaned = text

    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)

    # Remove most control characters, but preserve tabs/newlines.
    cleaned = re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", cleaned)
    cleaned = normalize_ocr_and_pdf_text(cleaned)

    return cleaned.strip()


# -----------------------------
# Formula / Scientific Token Protection
# -----------------------------
SUBSCRIPT_SUPERSCRIPT_CHARS = "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎"
SUB_SUP_PATTERN = f"[{re.escape(SUBSCRIPT_SUPERSCRIPT_CHARS)}]"


def protect_sensitive_tokens(text: str):
    """
    Protect tokens that translation engines often damage:
    - H₂O
    - CO₂
    - x²
    - 10⁻³
    - H2O / CO2 chemical-ish formulas

    This does not fix text that was extracted incorrectly from the PDF.
    It only prevents correctly extracted tokens from being translated/mangled.
    """
    if not text:
        return text, {}

    protected = {}

    patterns = [
        # Any contiguous token containing unicode subscript/superscript chars.
        rf"\b[\wΑ-Ωα-ωµ°+\-*/=().,]*{SUB_SUP_PATTERN}[\wΑ-Ωα-ωµ°+\-*/=().,]*\b",

        # Chemical-ish formulas with digits, e.g. H2O, CO2, Na2SO4.
        r"\b(?:[A-Z][a-z]?\d*){1,}[A-Z][a-z]?\d+\b",
        r"\b(?:[A-Z][a-z]?\d+){1,}(?:[A-Z][a-z]?\d*)*\b",
    ]

    combined = re.compile("|".join(f"({p})" for p in patterns))

    def replacement(match):
        token = match.group(0)

        # Avoid protecting pure numbers.
        if token.isdigit():
            return token

        placeholder = f"ZXQKEEP{len(protected)}QXZ"
        protected[placeholder] = token
        return placeholder

    protected_text = combined.sub(replacement, text)

    return protected_text, protected


def restore_sensitive_tokens(text: str, protected: dict):
    if not text or not protected:
        return text

    restored = text

    for placeholder, original in protected.items():
        restored = restored.replace(placeholder, original)

        # Sometimes translators add spaces around placeholder-like strings.
        spaced = " ".join(list(placeholder))
        restored = restored.replace(spaced, original)

    return restored


# -----------------------------
# Translation Helpers
# -----------------------------
def make_cache_key(text: str, source: str, target: str) -> str:
    digest = hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()
    return f"{source}:{target}:{digest}"


def split_text_for_translation(text: str, max_chars: int = 4300) -> list[str]:
    text = sanitize_text(text)

    if not text:
        return []

    if len(text) <= max_chars:
        return [text]

    pieces = re.split(r"(\n\s*\n)", text)

    chunks = []
    current = ""

    def flush_current():
        nonlocal current
        if current.strip():
            chunks.append(current.strip())
        current = ""

    for piece in pieces:
        if not piece.strip():
            continue

        if len(piece) > max_chars:
            flush_current()

            sentences = re.split(r"(?<=[.!?。！？])\s+", piece)
            sentence_buffer = ""

            for sentence in sentences:
                if len(sentence) > max_chars:
                    if sentence_buffer.strip():
                        chunks.append(sentence_buffer.strip())
                        sentence_buffer = ""

                    for i in range(0, len(sentence), max_chars):
                        chunks.append(sentence[i:i + max_chars].strip())

                else:
                    if len(sentence_buffer) + len(sentence) + 1 <= max_chars:
                        sentence_buffer += sentence + " "
                    else:
                        chunks.append(sentence_buffer.strip())
                        sentence_buffer = sentence + " "

            if sentence_buffer.strip():
                chunks.append(sentence_buffer.strip())

        else:
            if len(current) + len(piece) + 1 <= max_chars:
                current += piece + "\n"
            else:
                flush_current()
                current = piece + "\n"

    flush_current()

    return [chunk for chunk in chunks if chunk.strip()]


def translate_chunk(
    chunk,
    translator,
    source,
    target,
    retries=2,
    sleep_seconds=0.25,
    debug_context="",
):
    chunk = sanitize_text(chunk)

    if not chunk:
        return ""

    if source == target and source != "auto":
        return chunk

    cache_key = make_cache_key(chunk, source, target)

    if cache_key in st.session_state["translation_cache"]:
        return st.session_state["translation_cache"][cache_key]

    last_error = None

    for attempt in range(retries + 1):
        try:
            translated = translator.translate(chunk)

            if translated is None:
                raise ValueError("Translator returned None.")

            translated = sanitize_text(translated)

            st.session_state["translation_cache"][cache_key] = translated
            time.sleep(sleep_seconds)

            return translated

        except Exception as exc:
            last_error = exc
            time.sleep(0.75 * (attempt + 1))

    error_message = f"Translation failed at {debug_context}: {str(last_error)}"
    st.session_state["debug_log"].append(error_message)

    # Return original on failure to prevent data loss.
    return chunk


def translate_text(text, translator, source, target, debug_context=""):
    text = sanitize_text(text)

    if not text:
        return ""

    protected = {}

    if st.session_state.get("preserve_formula_tokens", True):
        text, protected = protect_sensitive_tokens(text)

    chunks = split_text_for_translation(text)

    translated_chunks = [
        translate_chunk(
            c,
            translator,
            source,
            target,
            debug_context=debug_context,
        )
        for c in chunks
    ]

    translated = "\n\n".join(translated_chunks)
    translated = restore_sensitive_tokens(translated, protected)
    translated = sanitize_text(translated)

    return translated


# -----------------------------
# PDF Font / Text Insertion Helpers
# -----------------------------
def page_insert_textbox_unicode(
    page,
    rect,
    text,
    fontsize,
    align=fitz.TEXT_ALIGN_LEFT,
    color=(0, 0, 0),
):
    """
    Insert text using a Unicode font when available.

    This is the main fix for characters like:
    - „
    - “
    - ”
    - ²
    - ₂
    - ⁻
    becoming '?' in normal PDF output.
    """
    text = sanitize_text(text)

    if UNICODE_FONT_PATH:
        return page.insert_textbox(
            rect,
            text,
            fontsize=fontsize,
            fontname=CUSTOM_FONT_NAME,
            fontfile=UNICODE_FONT_PATH,
            color=color,
            align=align,
        )

    # Fallback. Built-in Helvetica may not support all Unicode glyphs.
    return page.insert_textbox(
        rect,
        text,
        fontsize=fontsize,
        fontname="helv",
        color=color,
        align=align,
    )


def insert_textbox_autofit(
    page,
    rect: fitz.Rect,
    text: str,
    max_font_size=11.0,
    min_font_size=4.5,
):
    rect = fitz.Rect(rect)
    text = sanitize_text(text)

    # Padding
    rect.x0 += 1
    rect.y0 += 1
    rect.x1 -= 1
    rect.y1 -= 1

    if rect.width <= 2 or rect.height <= 2:
        return False

    font_size = max_font_size

    while font_size >= min_font_size:
        try:
            rc = page_insert_textbox_unicode(
                page,
                rect,
                text,
                fontsize=font_size,
            )

            if rc >= 0:
                return True

        except Exception as exc:
            st.session_state["debug_log"].append(
                f"Unicode PDF text insert failed: {str(exc)}"
            )

            try:
                rc = page.insert_textbox(
                    rect,
                    text,
                    fontsize=font_size,
                    fontname="helv",
                    color=(0, 0, 0),
                    align=fitz.TEXT_ALIGN_LEFT,
                )

                if rc >= 0:
                    return True

            except Exception as fallback_exc:
                st.session_state["debug_log"].append(
                    f"Fallback PDF text insert failed: {str(fallback_exc)}"
                )

        font_size -= 0.5

    # Fallback: truncate text.
    shortened = text

    while len(shortened) > 30:
        shortened = shortened[: int(len(shortened) * 0.9)].rstrip() + "..."

        try:
            rc = page_insert_textbox_unicode(
                page,
                rect,
                shortened,
                fontsize=min_font_size,
            )

            if rc >= 0:
                return False

        except Exception:
            pass

    return False


def extract_pdf_text_blocks(page):
    blocks = []

    page_dict = page.get_text("dict")

    for block in page_dict.get("blocks", []):
        if block.get("type") != 0 or "lines" not in block:
            continue

        lines = []

        for line in block["lines"]:
            spans = [
                sanitize_text(span.get("text", ""))
                for span in line.get("spans", [])
                if span.get("text", "").strip()
            ]

            line_text = "".join(spans).strip()

            if line_text:
                lines.append(line_text)

        block_text = "\n".join(lines).strip()
        block_text = sanitize_text(block_text)

        if not block_text:
            continue

        rect = fitz.Rect(block["bbox"])

        if rect.is_empty or rect.width < 4 or rect.height < 4:
            continue

        blocks.append((rect, block_text))

    return blocks


# -----------------------------
# OCR Helper Functions
# -----------------------------
def render_page_to_image(page, zoom=2.5):
    """
    Render a PDF page to a PIL image for OCR.
    """
    matrix = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=matrix, alpha=False)

    image = Image.frombytes(
        "RGB",
        [pix.width, pix.height],
        pix.samples,
    )

    return image


def ocr_page_text(page, source_code: str, zoom=2.5) -> str:
    """
    OCR a scanned PDF page using Tesseract.
    """
    image = render_page_to_image(page, zoom=zoom)

    tesseract_lang = TESSERACT_LANGUAGES.get(source_code, "eng")

    try:
        text = pytesseract.image_to_string(
            image,
            lang=tesseract_lang,
            config="--oem 1 --psm 6 -c preserve_interword_spaces=1",
        )

        return sanitize_text(text)

    except pytesseract.TesseractNotFoundError:
        raise RuntimeError(
            "Tesseract OCR was not found on the Streamlit server. "
            "Add tesseract-ocr to packages.txt and redeploy."
        )

    except Exception as exc:
        st.session_state["debug_log"].append(
            f"OCR failed: {str(exc)}"
        )
        return ""


def split_text_for_pdf_pages(text: str, max_chars: int = 2600) -> list[str]:
    """
    Split translated OCR text into page-sized chunks.
    """
    text = sanitize_text(text)

    if not text:
        return []

    paragraphs = re.split(r"\n\s*\n", text)

    pages = []
    current = ""

    for paragraph in paragraphs:
        paragraph = paragraph.strip()

        if not paragraph:
            continue

        if len(paragraph) > max_chars:
            if current.strip():
                pages.append(current.strip())
                current = ""

            for i in range(0, len(paragraph), max_chars):
                pages.append(paragraph[i:i + max_chars].strip())

        elif len(current) + len(paragraph) + 2 <= max_chars:
            current += paragraph + "\n\n"

        else:
            pages.append(current.strip())
            current = paragraph + "\n\n"

    if current.strip():
        pages.append(current.strip())

    return pages


def add_translated_ocr_pages(doc, translated_text: str, source_page_rect, title="OCR Translation"):
    """
    Add one or more new PDF pages containing translated OCR text.
    """
    page_chunks = split_text_for_pdf_pages(translated_text)

    if not page_chunks:
        return

    for chunk_index, chunk in enumerate(page_chunks):
        new_page = doc.new_page(
            width=source_page_rect.width,
            height=source_page_rect.height,
        )

        margin = 50

        title_rect = fitz.Rect(
            margin,
            25,
            source_page_rect.width - margin,
            55,
        )

        text_rect = fitz.Rect(
            margin,
            65,
            source_page_rect.width - margin,
            source_page_rect.height - margin,
        )

        heading = title

        if len(page_chunks) > 1:
            heading += f" — Part {chunk_index + 1}/{len(page_chunks)}"

        insert_textbox_autofit(
            new_page,
            title_rect,
            heading,
            max_font_size=12,
            min_font_size=8,
        )

        inserted = insert_textbox_autofit(
            new_page,
            text_rect,
            chunk,
            max_font_size=11,
            min_font_size=7,
        )

        if not inserted:
            st.session_state["debug_log"].append(
                "Some OCR translated text was truncated on an added OCR page."
            )


def overlay_translated_ocr_text(page, translated_text: str):
    """
    Experimental: cover the scanned page and insert translated OCR text.
    This will not preserve exact layout.
    """
    margin = 40

    rect = fitz.Rect(
        margin,
        margin,
        page.rect.width - margin,
        page.rect.height - margin,
    )

    page.draw_rect(
        rect,
        color=(1, 1, 1),
        fill=(1, 1, 1),
    )

    inserted = insert_textbox_autofit(
        page,
        rect,
        translated_text,
        max_font_size=10.5,
        min_font_size=5,
    )

    if not inserted:
        st.session_state["debug_log"].append(
            "OCR overlay text was truncated because it did not fit."
        )


# -----------------------------
# File Translation Functions
# -----------------------------
def translate_pdf(
    file_bytes,
    source,
    target,
    progress_callback=None,
    enable_ocr=False,
    ocr_output_mode="Append translated OCR pages",
    ocr_zoom_value=2.5,
):
    translator = GoogleTranslator(source=source, target=target)
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    original_page_count = len(doc)

    for page_index in range(original_page_count):
        page = doc[page_index]
        original_blocks = extract_pdf_text_blocks(page)

        # Case 1: PDF has selectable/embedded text.
        if original_blocks:
            translated_blocks = []

            for block_index, (rect, block_text) in enumerate(original_blocks):
                translated_text = translate_text(
                    block_text,
                    translator,
                    source,
                    target,
                    f"PDF Page {page_index + 1}, Block {block_index + 1}",
                )

                translated_blocks.append((rect, translated_text))

            # Redact original text.
            for rect, _ in translated_blocks:
                page.add_redact_annot(rect, fill=(1, 1, 1))

            if translated_blocks:
                page.apply_redactions(images=0)

            # Insert translated text using Unicode font.
            for rect, translated_text in translated_blocks:
                if not insert_textbox_autofit(page, rect, translated_text):
                    st.session_state["debug_log"].append(
                        f"Page {page_index + 1}: Text truncated due to limited space."
                    )

        # Case 2: Scanned/image-only PDF page.
        else:
            if enable_ocr:
                st.session_state["debug_log"].append(
                    f"Page {page_index + 1}: No embedded text found. Running OCR."
                )

                ocr_text = ocr_page_text(
                    page,
                    source,
                    zoom=ocr_zoom_value,
                )

                if not ocr_text:
                    st.session_state["debug_log"].append(
                        f"Page {page_index + 1}: OCR found no text."
                    )

                else:
                    translated_ocr_text = translate_text(
                        ocr_text,
                        translator,
                        source,
                        target,
                        f"OCR Page {page_index + 1}",
                    )

                    if ocr_output_mode == "Append translated OCR pages":
                        add_translated_ocr_pages(
                            doc,
                            translated_ocr_text,
                            page.rect,
                            title=f"Translated OCR text from page {page_index + 1}",
                        )

                    else:
                        overlay_translated_ocr_text(
                            page,
                            translated_ocr_text,
                        )

            else:
                st.session_state["debug_log"].append(
                    f"Page {page_index + 1}: No embedded text found. Enable OCR for scanned PDFs."
                )

        if progress_callback:
            progress_callback((page_index + 1) / original_page_count)

    output = io.BytesIO()

    doc.save(
        output,
        garbage=4,
        deflate=True,
    )

    doc.close()

    return output.getvalue()


def replace_paragraph_text(paragraph, new_text: str):
    if paragraph.runs:
        paragraph.runs[0].text = new_text

        for run in paragraph.runs[1:]:
            run.text = ""

    else:
        paragraph.add_run(new_text)


def iter_docx_paragraphs(document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph

        for paragraph in section.footer.paragraphs:
            yield paragraph

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph


def translate_docx(file_bytes, source, target, progress_callback=None):
    translator = GoogleTranslator(source=source, target=target)
    document = Document(io.BytesIO(file_bytes))

    paragraphs = [
        p for p in iter_docx_paragraphs(document)
        if p.text and p.text.strip()
    ]

    total = len(paragraphs)

    for index, paragraph in enumerate(paragraphs):
        translated_text = translate_text(
            paragraph.text,
            translator,
            source,
            target,
            f"DOCX Paragraph {index + 1}",
        )

        replace_paragraph_text(paragraph, translated_text)

        if progress_callback and total:
            progress_callback((index + 1) / total)

    output = io.BytesIO()
    document.save(output)

    return output.getvalue()


def translate_xlsx(file_bytes, source, target, progress_callback=None):
    translator = GoogleTranslator(source=source, target=target)
    workbook = load_workbook(io.BytesIO(file_bytes))

    text_cells = []

    for worksheet in workbook.worksheets:
        for row in worksheet.iter_rows():
            for cell in row:
                value = cell.value

                if isinstance(value, str) and value.strip() and not value.startswith("="):
                    text_cells.append(cell)

    total = len(text_cells)

    for index, cell in enumerate(text_cells):
        translated_text = translate_text(
            cell.value,
            translator,
            source,
            target,
            f"XLSX Cell {cell.coordinate}",
        )

        cell.value = translated_text

        if progress_callback and total:
            progress_callback((index + 1) / total)

    output = io.BytesIO()
    workbook.save(output)

    return output.getvalue()


def get_mime_type(ext: str) -> str:
    return {
        "pdf": "application/pdf",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    }.get(ext, "application/octet-stream")


# -----------------------------
# Main UI
# -----------------------------
mode = st.radio(
    "Select Mode",
    ["Upload Files", "Plain Text"],
    horizontal=True,
)

if src_code == tgt_code and src_code != "auto":
    st.warning("Source and target languages are the same.")


if mode == "Upload Files":
    st.subheader("Upload PDF, Word, or Excel File")

    uploaded_file = st.file_uploader(
        "Choose a file",
        type=["pdf", "docx", "xlsx"],
        label_visibility="collapsed",
    )

    if uploaded_file:
        st.success(f"Loaded: **{uploaded_file.name}**")

        if st.button("Process & Translate", type="primary", use_container_width=True):
            st.session_state["debug_log"] = []
            st.session_state["processed_output_data"] = None
            st.session_state["processed_output_name"] = ""

            progress_bar = st.progress(0)
            status_text = st.empty()

            try:
                file_bytes = uploaded_file.read()
                file_name = uploaded_file.name
                extension = file_name.split(".")[-1].lower()

                def update_progress(val):
                    progress_bar.progress(min(max(val, 0), 1))

                status_text.text("Translating...")

                if extension == "pdf":
                    output_bytes = translate_pdf(
                        file_bytes,
                        src_code,
                        tgt_code,
                        progress_callback=update_progress,
                        enable_ocr=enable_ocr,
                        ocr_output_mode=ocr_output_mode,
                        ocr_zoom_value=ocr_zoom,
                    )

                elif extension == "docx":
                    output_bytes = translate_docx(
                        file_bytes,
                        src_code,
                        tgt_code,
                        update_progress,
                    )

                elif extension == "xlsx":
                    output_bytes = translate_xlsx(
                        file_bytes,
                        src_code,
                        tgt_code,
                        update_progress,
                    )

                else:
                    st.error("Unsupported file type.")
                    st.stop()

                st.session_state["processed_output_data"] = output_bytes
                st.session_state["processed_output_name"] = f"translated_{file_name}"

                progress_bar.progress(1.0)
                status_text.text("Done.")

                st.success("Translation complete!")

                if debug_mode and st.session_state["debug_log"]:
                    with st.expander("View Error Log"):
                        for log_entry in st.session_state["debug_log"]:
                            st.code(log_entry)

            except Exception as exc:
                st.error(f"Error: {str(exc)}")

                if debug_mode:
                    st.exception(exc)

    if st.session_state["processed_output_data"]:
        st.divider()
        st.subheader("Download Result")

        extension = st.session_state["processed_output_name"].split(".")[-1].lower()

        st.download_button(
            label="📥 Download Translated File",
            data=st.session_state["processed_output_data"],
            file_name=st.session_state["processed_output_name"],
            mime=get_mime_type(extension),
            use_container_width=True,
        )


elif mode == "Plain Text":
    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader("Input")

        input_text = st.text_area(
            "Paste text here",
            height=400,
            placeholder="Type or paste...",
            label_visibility="collapsed",
        )

        translate_button = st.button(
            "Translate Text",
            type="primary",
            use_container_width=True,
        )

    with col2:
        st.subheader("Translation")

        if translate_button:
            st.session_state["debug_log"] = []

            if not input_text.strip():
                st.warning("Please enter text.")

            else:
                try:
                    translator = GoogleTranslator(source=src_code, target=tgt_code)

                    st.session_state["translated_text_result"] = translate_text(
                        input_text,
                        translator,
                        src_code,
                        tgt_code,
                        "Plain Text",
                    )

                except Exception as exc:
                    st.session_state["translated_text_result"] = f"Error: {str(exc)}"

        if st.session_state["translated_text_result"]:
            st.text_area(
                "Output",
                value=st.session_state["translated_text_result"],
                height=400,
                label_visibility="collapsed",
            )

            st.download_button(
                "Download .txt",
                data=st.session_state["translated_text_result"],
                file_name="translated.txt",
                mime="text/plain",
                use_container_width=True,
            )

        else:
            st.info("Translation will appear here.")
